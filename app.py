import io
import os
import zipfile
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
import pypdfium2 as pdfium
from PIL import Image
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas
import streamlit as st

# Configure Streamlit page layout
st.set_page_config(
    page_title="SUST EEE Cover Page Generator", page_icon="📄", layout="centered"
)

# -----------------------------------------------------------------------------
# 0. DATABASE & CONSTANTS
# -----------------------------------------------------------------------------
SUST_EEE_TEACHERS = {
    "Select Teacher": {"designation": ""},
    "Dr. Ifte Khairul Amin": {"designation": "Professor"},
    "Biswajit Paul": {"designation": "Associate Professor"},
    "Dr. Md Rasedujjaman": {"designation": "Associate Professor"},
    "Dr. Mohammad Kamruzzaman Khan Prince": {"designation": "Associate Professor"},
    "Arif Ahammad": {"designation": "Assistant Professor"},
    "Hriteshwar Talukder": {"designation": "Assistant Professor"},
    "Tahmid Aziz Chowdhury": {"designation": "Assistant Professor"},
    "Md. Asaduz Zaman Mamun": {"designation": "Assistant Professor"},
    "Showmik Singha": {"designation": "Assistant Professor"},
    "Nafis Imtiaz Rahman": {"designation": "Assistant Professor"},
    "Md. Ishfak Tahmid": {"designation": "Assistant Professor"},
    "Tuhin Dey": {"designation": "Assistant Professor"},
    "Md. Shariful Islam": {"designation": "Lecturer"},
    "Mohona Das Gupta": {"designation": "Lecturer"},
    "Naima Sultana Alam Supti": {"designation": "Lecturer"},
    "Md. Kabir Hasan": {"designation": "Lecturer"},
    "Afshana Begum": {"designation": "Lecturer"},
    "Md. Tasfiq Rahman": {"designation": "Lecturer"},
    "Myen Uddin": {"designation": "Lecturer"},
    "Irfan Nafiz Shahan": {"designation": "Lecturer"},
}

DESIGNATION_OPTIONS = [
    "Professor",
    "Associate Professor",
    "Assistant Professor",
    "Lecturer",
    "Others",
]

SUST_DEPARTMENTS = [
    "Electrical and Electronic Engineering",
    "Computer Science and Engineering",
    "Chemical Engineering and Polymer Science",
    "Civil and Environmental Engineering",
    "Industrial and Production Engineering",
    "Mechanical Engineering",
    "Petroleum and Mining Engineering",
    "Software Engineering",
    "Food Engineering and Tea Technology",
    "Architecture",
    "Physics",
    "Chemistry",
    "Mathematics",
    "Statistics",
    "Geography and Environment",
    "Oceanography",
    "Biochemistry and Molecular Biology",
    "Genetic Engineering and Biotechnology",
    "Business Administration",
    "Economics",
    "Social Work",
    "Sociology",
    "Political Studies",
    "Public Administration",
    "Anthropology",
    "English",
    "Bangla",
    "Institute of Information and Communication Technology (IICT)",
    "Others",
]

# -----------------------------------------------------------------------------
# 1. SESSION STATE & INITIALIZATION
# -----------------------------------------------------------------------------
if "step" not in st.session_state:
    st.session_state.step = "form"

if "form_version" not in st.session_state:
    st.session_state.form_version = 0

DEFAULT_DATA = {
    "doc_type_choice": "Lab Report",
    "custom_doc_type": "",
    "work_type": "Individual",
    "exp_no": "",
    "title": "",
    "course_title": "",
    "course_code": "",
    "student_name": "",
    "reg_no": "",
    "group_no": "",
    "num_students": 2,
    "teacher_name": "",
    "designation": "",
    "department": "Electrical and Electronic Engineering",
    "custom_dept": "",
    "sub_date": "",
    "is_external_teacher": False,
}

if "form_data" not in st.session_state:
    st.session_state.form_data = DEFAULT_DATA.copy()


def clear_all_form_fields():
    st.session_state.form_data = DEFAULT_DATA.copy()
    st.session_state.form_version += 1


def prepare_image(image_path):
    if not os.path.exists(image_path):
        return None
    try:
        img = Image.open(image_path)
        return ImageReader(img)
    except Exception:
        return None


# -----------------------------------------------------------------------------
# 2. WARNING MODAL FOR BLANK FIELDS
# -----------------------------------------------------------------------------
@st.dialog("⚠️ Missing Information Warning")
def show_missing_fields_warning(missing_fields):
    st.write(
        "The following fields are currently **blank**. If you proceed, they will remain empty on the generated documents so you can write them in by hand."
    )
    for field in missing_fields:
        st.markdown(f"- **{field}**")

    st.write("Would you like to continue to the preview or edit your details?")

    col1, col2 = st.columns(2)
    with col1:
        if st.button("Proceed Anyway", type="primary", use_container_width=True):
            st.session_state.step = "preview"
            st.rerun()
    with col2:
        if st.button("Go Back & Fill", use_container_width=True):
            st.rerun()


# -----------------------------------------------------------------------------
# 3. PDF GENERATION LOGIC
# -----------------------------------------------------------------------------
def generate_pdf():
    data = st.session_state.form_data
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    # --- 1. Document Type Header ---
    doc_type_text = data.get("doc_type_choice", "Lab Report")
    if doc_type_text == "Others":
        doc_type_text = data.get("custom_doc_type", "Cover Page")

    current_y = height - 80
    c.setFont("Times-Bold", 22)
    c.drawCentredString(width / 2, current_y, doc_type_text.upper())

    # --- 2. Top Horizontal Line ---
    current_y -= 25
    c.setLineWidth(1)
    c.line(50, current_y, width - 50, current_y)

    # --- 3. Exp / Assignment / Project / Proposal Number & Title ---
    exp_num = data.get("exp_no", "").strip()
    if exp_num:
        current_y -= 25
        c.setFont("Times-Roman", 14)
        if doc_type_text == "Lab Report":
            label_str = f"Part of Exp : {exp_num}"
        elif doc_type_text == "Assignment":
            label_str = f"Assignment No : {exp_num}"
        elif doc_type_text == "Project Report":
            label_str = f"Project No : {exp_num}"
        elif doc_type_text == "Project Proposal":
            label_str = f"Proposal No : {exp_num}"
        else:
            label_str = f"No : {exp_num}"

        c.drawCentredString(width / 2, current_y, label_str)

    if data.get("title"):
        current_y -= 25
        c.setFont("Times-Bold", 16)
        c.drawCentredString(width / 2, current_y, data["title"])

    # --- 4. Bottom Horizontal Line ---
    current_y -= 15
    c.line(50, current_y, width - 50, current_y)

    # --- 5. Centered EEE Department Logo ---
    current_y -= 100
    logo_w, logo_h = 90, 90

    eee_logo = prepare_image("logo_eee.png") or prepare_image(
        "eee-sust-logo-png_seeklogo-535291.png"
    )
    if eee_logo:
        c.drawImage(
            eee_logo,
            (width - logo_w) / 2,
            current_y,
            width=logo_w,
            height=logo_h,
            mask="auto",
        )

    # --- 6. Department Name & Course Info ---
    current_y -= 45
    c.setFont("Times-Bold", 15)
    c.drawCentredString(
        width / 2,
        current_y,
        "Department of Electrical and Electronic Engineering",
    )

    if data.get("course_title"):
        current_y -= 35
        c.setFont("Times-Roman", 13)
        c.drawCentredString(
            width / 2,
            current_y,
            f"Course Title: {data['course_title']}",
        )

    if data.get("course_code"):
        current_y -= 25
        c.setFont("Times-Roman", 13)
        c.drawCentredString(
            width / 2,
            current_y,
            f"Course Code: {data['course_code']}",
        )

    # --- 7. Submitted By Section ---
    current_y -= 55
    c.setFont("Times-Bold", 13)
    c.drawCentredString(width / 2, current_y, "Submitted By:")
    c.line(width / 2 - 45, current_y - 3, width / 2 + 45, current_y - 3)

    c.setFont("Times-Roman", 12)
    current_y -= 25

    if data.get("work_type") == "Individual":
        if data.get("student_name"):
            c.drawCentredString(width / 2, current_y, data["student_name"])
            current_y -= 20
        if data.get("reg_no"):
            c.drawCentredString(
                width / 2,
                current_y,
                f"Registration No: {data['reg_no']}",
            )
            current_y -= 20
    else:
        for i in range(int(data.get("num_students", 2))):
            s_name = data.get(f"sname_{i}", "")
            s_reg = data.get(f"sreg_{i}", "")
            if s_name:
                info = f"{s_name} (Reg: {s_reg})" if s_reg else s_name
                c.drawCentredString(width / 2, current_y, info)
                current_y -= 18

        if data.get("group_no"):
            current_y -= 5
            c.drawCentredString(
                width / 2, current_y, f"Group No : {data['group_no']}"
            )
            current_y -= 20

    # --- 8. Submitted To Section ---
    current_y -= 35
    c.setFont("Times-Bold", 13)
    c.drawCentredString(width / 2, current_y, "Submitted To:")
    c.line(width / 2 - 45, current_y - 3, width / 2 + 45, current_y - 3)

    c.setFont("Times-Roman", 12)
    current_y -= 25

    if data.get("teacher_name"):
        c.drawCentredString(width / 2, current_y, data["teacher_name"])
        current_y -= 20
    if data.get("designation"):
        c.drawCentredString(width / 2, current_y, data["designation"])
        current_y -= 20

    # Resolve Department
    if data.get("is_external_teacher"):
        dept_val = data.get("department", "Electrical and Electronic Engineering")
        if dept_val == "Others":
            dept_val = data.get("custom_dept", "Department")
        dept_str = f"Department of {dept_val}" if not dept_val.lower().startswith("department") else dept_val
    else:
        dept_str = "Department of Electrical & Electronic Engineering"

    c.drawCentredString(width / 2, current_y, dept_str)
    current_y -= 20
    c.drawCentredString(
        width / 2,
        current_y,
        "Shahjalal University of Science and Technology, Sylhet",
    )

    # --- 9. Date of Submission ---
    sub_date = data.get("sub_date", "").strip()
    c.setFont("Times-Roman", 12)
    if sub_date:
        c.drawCentredString(width / 2, 50, f"Date of Submission: {sub_date}")
    else:
        c.drawCentredString(width / 2, 50, "Date of Submission: ....................................")

    c.showPage()
    c.save()
    buffer.seek(0)
    return buffer


# -----------------------------------------------------------------------------
# 4. DOCX GENERATION LOGIC (NATIVE MS WORD)
# -----------------------------------------------------------------------------
def generate_docx():
    data = st.session_state.form_data
    doc = Document()

    # Set page margins to 1 inch
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    def add_centered_p(text="", bold=False, size=12, space_after=6, underline=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(space_after)
        if text:
            run = p.add_run(text)
            run.bold = bold
            run.underline = underline
            run.font.name = "Times New Roman"
            run.font.size = Pt(size)
        return p

    # --- 1. Document Type Header ---
    doc_type_text = data.get("doc_type_choice", "Lab Report")
    if doc_type_text == "Others":
        doc_type_text = data.get("custom_doc_type", "Cover Page")

    add_centered_p(doc_type_text.upper(), bold=True, size=22, space_after=12)

    # --- 2. Exp / Assignment / Proposal Number & Title ---
    exp_num = data.get("exp_no", "").strip()
    if exp_num:
        if doc_type_text == "Lab Report":
            label_str = f"Part of Exp : {exp_num}"
        elif doc_type_text == "Assignment":
            label_str = f"Assignment No : {exp_num}"
        elif doc_type_text == "Project Report":
            label_str = f"Project No : {exp_num}"
        elif doc_type_text == "Project Proposal":
            label_str = f"Proposal No : {exp_num}"
        else:
            label_str = f"No : {exp_num}"

        add_centered_p(label_str, size=14, space_after=6)

    if data.get("title"):
        add_centered_p(data["title"], bold=True, size=16, space_after=18)

    # --- 3. Logo ---
    eee_path = (
        "logo_eee.png"
        if os.path.exists("logo_eee.png")
        else "eee-sust-logo-png_seeklogo-535291.png"
    )
    if os.path.exists(eee_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(12)
        p_img.paragraph_format.space_after = Pt(18)
        p_img.add_run().add_picture(eee_path, width=Inches(1.2))

    # --- 4. Department Name & Course Info ---
    add_centered_p(
        "Department of Electrical and Electronic Engineering",
        bold=True,
        size=15,
        space_after=12,
    )

    if data.get("course_title"):
        add_centered_p(f"Course Title: {data['course_title']}", size=13, space_after=4)

    if data.get("course_code"):
        add_centered_p(f"Course Code: {data['course_code']}", size=13, space_after=24)

    # --- 5. Submitted By Section ---
    add_centered_p("Submitted By:", bold=True, size=13, space_after=6, underline=True)

    if data.get("work_type") == "Individual":
        if data.get("student_name"):
            add_centered_p(data["student_name"], size=12, space_after=4)
        if data.get("reg_no"):
            add_centered_p(f"Registration No: {data['reg_no']}", size=12, space_after=20)
    else:
        for i in range(int(data.get("num_students", 2))):
            s_name = data.get(f"sname_{i}", "")
            s_reg = data.get(f"sreg_{i}", "")
            if s_name:
                info = f"{s_name} (Reg: {s_reg})" if s_reg else s_name
                add_centered_p(info, size=12, space_after=4)

        if data.get("group_no"):
            add_centered_p(f"Group No : {data['group_no']}", size=12, space_after=20)

    # --- 6. Submitted To Section ---
    add_centered_p("Submitted To:", bold=True, size=13, space_after=6, underline=True)

    if data.get("teacher_name"):
        add_centered_p(data["teacher_name"], size=12, space_after=4)
    if data.get("designation"):
        add_centered_p(data["designation"], size=12, space_after=4)

    # Resolve Department
    if data.get("is_external_teacher"):
        dept_val = data.get("department", "Electrical and Electronic Engineering")
        if dept_val == "Others":
            dept_val = data.get("custom_dept", "Department")
        dept_str = f"Department of {dept_val}" if not dept_val.lower().startswith("department") else dept_val
    else:
        dept_str = "Department of Electrical & Electronic Engineering"

    add_centered_p(dept_str, size=12, space_after=4)
    add_centered_p(
        "Shahjalal University of Science and Technology, Sylhet",
        size=12,
        space_after=30,
    )

    # --- 7. Date of Submission ---
    sub_date = data.get("sub_date", "").strip()
    date_text = (
        f"Date of Submission: {sub_date}"
        if sub_date
        else "Date of Submission: ...................................."
    )
    add_centered_p(date_text, size=12, space_after=0)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# -----------------------------------------------------------------------------
# 5. LATEX GENERATION & ZIP BUNDLING LOGIC
# -----------------------------------------------------------------------------
def generate_tex():
    data = st.session_state.form_data

    # Resolve document type
    doc_type = data.get("doc_type_choice", "Lab Report")
    if doc_type == "Others":
        doc_type = data.get("custom_doc_type", "Cover Page")

    exp_num = data.get("exp_no", "").strip()
    if exp_num:
        if doc_type == "Lab Report":
            num_str = f"Part of Exp : {exp_num}"
        elif doc_type == "Assignment":
            num_str = f"Assignment No : {exp_num}"
        elif doc_type == "Project Report":
            num_str = f"Project No : {exp_num}"
        elif doc_type == "Project Proposal":
            num_str = f"Proposal No : {exp_num}"
        else:
            num_str = f"No : {exp_num}"
    else:
        num_str = ""

    # Build student block
    student_block = ""
    if data.get("work_type") == "Individual":
        if data.get("student_name"):
            student_block += f"{data['student_name']}\\\\\n"
        if data.get("reg_no"):
            student_block += f"Registration No: {data['reg_no']}\\\\\n"
    else:
        for i in range(int(data.get("num_students", 2))):
            s_name = data.get(f"sname_{i}", "")
            s_reg = data.get(f"sreg_{i}", "")
            if s_name:
                info = f"{s_name} (Reg: {s_reg})" if s_reg else s_name
                student_block += f"{info}\\\\\n"
        if data.get("group_no"):
            student_block += f"Group No : {data['group_no']}\\\\\n"

    # Build teacher department string
    if data.get("is_external_teacher"):
        dept_val = data.get("department", "Electrical and Electronic Engineering")
        if dept_val == "Others":
            dept_val = data.get("custom_dept", "Department")
        dept_str = (
            f"Department of {dept_val}"
            if not dept_val.lower().startswith("department")
            else dept_val
        )
    else:
        dept_str = "Department of Electrical \\& Electronic Engineering"

    sub_date = data.get("sub_date", "").strip()
    date_str = (
        f"Date of Submission: {sub_date}"
        if sub_date
        else "Date of Submission: \\dots\\dots\\dots\\dots\\dots\\dots\\dots\\dots"
    )

    # Template
    tex_code = f"""\\documentclass[12pt,a4paper]{{article}}
\\usepackage[margin=1in]{{geometry}}
\\usepackage{{graphicx}}
\\usepackage{{setspace}}
\\pagestyle{{empty}}

\\begin{{document}}
\\begin{{center}}

{{\\Huge \\textbf{{{doc_type.upper()}}}}} \\\\[1.5em]
\\hrule height 1pt
\\vspace{{1em}}

\\Large {num_str} \\\\[0.8em]
{{\\Large \\textbf{{{data.get("title", "")}}}}} \\\\[1em]

\\hrule height 1pt
\\vspace{{2em}}

\\IfFileExists{{logo_eee.png}}{{
    \\includegraphics[width=1.2in]{{logo_eee.png}} \\\\[1.5em]
}}{{
    \\vspace{{1.2in}}
}}

{{\\large \\textbf{{Department of Electrical and Electronic Engineering}}}} \\\\[0.8em]
Course Title: {data.get("course_title", "")} \\\\[0.5em]
Course Code: {data.get("course_code", "")} \\\\[2em]

{{\\large \\underline{{\\textbf{{Submitted By:}}}}}} \\\\[0.8em]
{student_block}
\\vspace{{1.5em}}

{{\\large \\underline{{\\textbf{{Submitted To:}}}}}} \\\\[0.8em]
{data.get("teacher_name", "")} \\\\[0.4em]
{data.get("designation", "")} \\\\[0.4em]
{dept_str} \\\\[0.4em]
Shahjalal University of Science and Technology, Sylhet \\\\[3em]

\\vfill
{date_str}

\\end{{center}}
\\end{{document}}
"""
    return tex_code


def generate_tex_zip(tex_code):
    zip_buffer = io.BytesIO()

    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        # Add the generated .tex file
        zip_file.writestr("Cover_Page.tex", tex_code)

        # Include the EEE logo if available
        eee_path = (
            "logo_eee.png"
            if os.path.exists("logo_eee.png")
            else "eee-sust-logo-png_seeklogo-535291.png"
        )
        if os.path.exists(eee_path):
            zip_file.write(eee_path, arcname="logo_eee.png")

    zip_buffer.seek(0)
    return zip_buffer


def get_pdf_preview(pdf_bytes):
    pdf = pdfium.PdfDocument(pdf_bytes)
    page = pdf[0]
    image = page.render(scale=2).to_pil()
    return image


# -----------------------------------------------------------------------------
# 6. STREAMLIT UI LAYOUT
# -----------------------------------------------------------------------------
if st.session_state.step == "form":
    eee_path = (
        "logo_eee.png"
        if os.path.exists("logo_eee.png")
        else "eee-sust-logo-png_seeklogo-535291.png"
    )
    if os.path.exists(eee_path):
        st.image(eee_path, width=90)

    st.title("📄 SUST EEE Cover Page Generator")

    doc_options = ["Lab Report", "Assignment", "Project Report", "Project Proposal", "Others"]
    work_options = ["Individual", "Group Work"]

    v = st.session_state.form_version
    d = st.session_state.form_data

    # 1. Select Cover Page Type
    current_choice = d.get("doc_type_choice", "Lab Report")
    idx_doc = doc_options.index(current_choice) if current_choice in doc_options else 0

    d["doc_type_choice"] = st.selectbox(
        "Select Cover Page Type",
        doc_options,
        index=idx_doc,
        key=f"doc_type_{v}",
    )

    if d["doc_type_choice"] == "Others":
        d["custom_doc_type"] = st.text_input(
            "Specify Cover Page Type",
            value=d.get("custom_doc_type", ""),
            key=f"custom_doc_{v}",
        )

    st.subheader("Cover Page Details")

    # 2. Number / ID, Title, Course Name, Course Code
    exp_field_label = "Experiment No (e.g., 02)"
    if d["doc_type_choice"] == "Assignment":
        exp_field_label = "Assignment No (e.g., 01)"
    elif d["doc_type_choice"] == "Project Report":
        exp_field_label = "Project No (e.g., 01)"
    elif d["doc_type_choice"] == "Project Proposal":
        exp_field_label = "Proposal No (e.g., 01)"
    elif d["doc_type_choice"] == "Others":
        exp_field_label = "Number / ID (e.g., 01)"

    d["exp_no"] = st.text_input(
        exp_field_label,
        value=d.get("exp_no", ""),
        key=f"exp_no_{v}",
    )
    d["title"] = st.text_input(
        "Title",
        value=d.get("title", ""),
        key=f"title_{v}",
    )
    d["course_title"] = st.text_input(
        "Course Title",
        value=d.get("course_title", ""),
        key=f"course_title_{v}",
    )
    d["course_code"] = st.text_input(
        "Course Code",
        value=d.get("course_code", ""),
        key=f"course_code_{v}",
    )

    # 3. Individual or Group Option
    st.subheader("Student Details")
    d["work_type"] = st.radio(
        "Select Submission Type",
        work_options,
        index=work_options.index(d.get("work_type", "Individual")),
        horizontal=True,
        key=f"work_type_{v}",
    )

    # 4. Student Information
    if d["work_type"] == "Individual":
        d["student_name"] = st.text_input(
            "Student Name",
            value=d.get("student_name", ""),
            key=f"student_name_{v}",
        )
        d["reg_no"] = st.text_input(
            "Registration No",
            value=d.get("reg_no", ""),
            key=f"reg_no_{v}",
        )
    else:
        d["group_no"] = st.text_input(
            "Group No",
            value=d.get("group_no", ""),
            key=f"group_no_{v}",
        )
        d["num_students"] = st.number_input(
            "Number of Students",
            min_value=2,
            max_value=10,
            value=int(d.get("num_students", 2)),
            key=f"num_students_{v}",
        )

        for i in range(int(d["num_students"])):
            col1, col2 = st.columns(2)
            d[f"sname_{i}"] = col1.text_input(
                f"Student {i+1} Name",
                value=d.get(f"sname_{i}", ""),
                key=f"sname_{i}_{v}",
            )
            d[f"sreg_{i}"] = col2.text_input(
                f"Student {i+1} Reg No",
                value=d.get(f"sreg_{i}", ""),
                key=f"sreg_{i}_{v}",
            )

    # 5. Teacher Details Section
    st.subheader("Teacher Details")

    col_t1, col_t2 = st.columns([2, 1])
    with col_t2:
        is_external = st.checkbox(
            "Teacher from another department?",
            value=d.get("is_external_teacher", False),
            key=f"ext_teacher_{v}",
        )
        d["is_external_teacher"] = is_external

    with col_t1:
        teacher_list = list(SUST_EEE_TEACHERS.keys())
        selected_teacher_idx = 0
        if d.get("teacher_name") in teacher_list:
            selected_teacher_idx = teacher_list.index(d["teacher_name"])

        selected_teacher = st.selectbox(
            "Select SUST EEE Teacher",
            teacher_list,
            index=selected_teacher_idx,
            disabled=is_external,
            key=f"teacher_select_{v}",
        )

    if not is_external:
        if selected_teacher != "Select Teacher":
            d["teacher_name"] = selected_teacher
            d["designation"] = SUST_EEE_TEACHERS[selected_teacher]["designation"]
        else:
            d["teacher_name"] = ""
            d["designation"] = ""

        d["department"] = "Electrical and Electronic Engineering"

        st.text_input(
            "Teacher Name",
            value=d.get("teacher_name", ""),
            disabled=True,
            key=f"tname_dis_{v}",
        )
        st.text_input(
            "Teacher Designation",
            value=d.get("designation", ""),
            disabled=True,
            key=f"tdesig_dis_{v}",
        )
        st.text_input(
            "Department",
            value=d.get("department", ""),
            disabled=True,
            key=f"tdept_dis_{v}",
        )
    else:
        d["teacher_name"] = st.text_input(
            "Teacher Name",
            value=d.get("teacher_name", ""),
            key=f"tname_ext_{v}",
        )

        curr_desig = d.get("designation", "Assistant Professor")
        desig_idx = (
            DESIGNATION_OPTIONS.index(curr_desig)
            if curr_desig in DESIGNATION_OPTIONS
            else 4
        )

        selected_desig = st.selectbox(
            "Teacher Designation",
            DESIGNATION_OPTIONS,
            index=desig_idx,
            key=f"desig_select_{v}",
        )

        if selected_desig == "Others":
            d["designation"] = st.text_input(
                "Specify Designation",
                value=curr_desig if curr_desig not in DESIGNATION_OPTIONS else "",
                key=f"custom_desig_{v}",
            )
        else:
            d["designation"] = selected_desig

        curr_dept = d.get("department", "Electrical and Electronic Engineering")
        dept_idx = (
            SUST_DEPARTMENTS.index(curr_dept)
            if curr_dept in SUST_DEPARTMENTS
            else (len(SUST_DEPARTMENTS) - 1)
        )

        selected_dept = st.selectbox(
            "Teacher Department (Type to Search)",
            SUST_DEPARTMENTS,
            index=dept_idx,
            key=f"dept_select_{v}",
        )

        if selected_dept == "Others":
            d["department"] = "Others"
            d["custom_dept"] = st.text_input(
                "Specify Department Name",
                value=d.get("custom_dept", ""),
                key=f"custom_dept_{v}",
            )
        else:
            d["department"] = selected_dept

    d["sub_date"] = st.text_input(
        "Date of Submission (Leave blank for printable dotted line)",
        value=d.get("sub_date", ""),
        key=f"sub_date_{v}",
    )

    st.session_state.form_data = d

    action_col1, action_col2 = st.columns(2)

    with action_col1:
        if st.button("Generate & Preview", type="primary", use_container_width=True):
            missing_fields = []
            if not d.get("exp_no", "").strip():
                missing_fields.append(exp_field_label)
            if not d.get("title", "").strip():
                missing_fields.append("Title")
            if not d.get("course_title", "").strip():
                missing_fields.append("Course Title")
            if not d.get("course_code", "").strip():
                missing_fields.append("Course Code")

            if d["work_type"] == "Individual":
                if not d.get("student_name", "").strip():
                    missing_fields.append("Student Name")
                if not d.get("reg_no", "").strip():
                    missing_fields.append("Registration No")
            else:
                if not d.get("group_no", "").strip():
                    missing_fields.append("Group No")

            if not d.get("teacher_name", "").strip():
                missing_fields.append("Teacher Name")

            if d.get("is_external_teacher"):
                if d.get("department") == "Others" and not d.get("custom_dept", "").strip():
                    missing_fields.append("Specified Department")

            if missing_fields:
                show_missing_fields_warning(missing_fields)
            else:
                st.session_state.step = "preview"
                st.rerun()

    with action_col2:
        if st.button("🗑️ Clear Form", use_container_width=True):
            clear_all_form_fields()
            st.rerun()

elif st.session_state.step == "preview":
    st.title("📄 SUST EEE Cover Page Generator")
    st.subheader("Preview Your Cover Page")

    pdf_buffer = generate_pdf()
    pdf_bytes = pdf_buffer.getvalue()

    docx_buffer = generate_docx()
    docx_bytes = docx_buffer.getvalue()

    tex_code = generate_tex()
    tex_zip_buffer = generate_tex_zip(tex_code)
    tex_zip_bytes = tex_zip_buffer.getvalue()

    preview_img = get_pdf_preview(pdf_bytes)
    st.image(preview_img, use_container_width=True)

    btn_col1, btn_col2, btn_col3, btn_col4 = st.columns([1, 1, 1, 1.2])

    with btn_col1:
        if st.button("✏️ Edit Details", use_container_width=True):
            st.session_state.step = "form"
            st.rerun()

    with btn_col2:
        st.download_button(
            label="📥 PDF",
            data=pdf_bytes,
            file_name="Cover_Page.pdf",
            mime="application/pdf",
            type="primary",
            use_container_width=True,
        )

    with btn_col3:
        st.download_button(
            label="📝 DOCX",
            data=docx_bytes,
            file_name="Cover_Page.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

    with btn_col4:
        st.download_button(
            label="📦 LaTeX Bundle (.zip)",
            data=tex_zip_bytes,
            file_name="LaTeX_Cover_Page.zip",
            mime="application/zip",
            use_container_width=True,
        )
