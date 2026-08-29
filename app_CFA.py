import io
import os
import zipfile
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
import pypdfium2 as pdfium
from PIL import Image
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas
from reportlab.platypus import Paragraph
import streamlit as st

# Configure Streamlit page layout
st.set_page_config(
    page_title="Cover Page Generator", page_icon="📄", layout="centered"
)

# -----------------------------------------------------------------------------
# 0. CONSTANTS & INSTITUTION LISTS
# -----------------------------------------------------------------------------
BANGLADESH_UNIVERSITIES = [
    # --- Top World Reputed Universities ---
    "Harvard University",
    "Massachusetts Institute of Technology (MIT)",
    "Stanford University",
    "University of Cambridge",
    "University of Oxford",
    "Imperial College London",
    "ETH Zurich",
    "California Institute of Technology (Caltech)",
    "National University of Singapore (NUS)",
    "University College London (UCL)",
    "Nanyang Technological University (NTU)",
    "Princeton University",
    "Yale University",
    "Cornell University",
    "Columbia University",
    "University of California, Berkeley (UCB)",
    "University of Chicago",
    "Tsinghua University",
    "Peking University",
    "University of Toronto",
    "The University of Melbourne",
    "The University of Sydney",
    "The University of Tokyo",
    "Seoul National University",

    # --- Public Universities (Bangladesh) ---
    "University of Dhaka (DU)",
    "Bangladesh University of Engineering & Technology (BUET)",
    "University of Chittagong (CU)",
    "University of Rajshahi (RU)",
    "Jahangirnagar University (JU)",
    "Shahjalal University of Science & Technology (SUST)",
    "Bangladesh Agricultural University (BAU)",
    "Khulna University (KU)",
    "Islamic University, Bangladesh (IU)",
    "Jagannath University (JnU)",
    "Comilla University (CoU)",
    "Chittagong University of Engineering & Technology (CUET)",
    "Rajshahi University of Engineering & Technology (RUET)",
    "Khulna University of Engineering & Technology (KUET)",
    "Dhaka University of Engineering & Technology (DUET)",
    "Noakhali Science & Technology University (NSTU)",
    "Hajee Mohammad Danesh Science & Technology University (HSTU)",
    "Mawlana Bhashani Science & Technology University (MBSTU)",
    "Patuakhali Science & Technology University (PSTU)",
    "Sher-e-Bangla Agricultural University (SAU)",
    "Jessore University of Science & Technology (JUST)",
    "Pabna University of Science & Technology (PUST)",
    "Begum Rokeya University, Rangpur (BRUR)",
    "Bangladesh University of Professionals (BUP)",
    "Bangladesh University of Textiles (BUTEX)",
    "University of Barishal (BU)",
    "Jatiya Kabi Kazi Nazrul Islam University (JKKNIU)",
    "Sylhet Agricultural University (SAU)",
    "Chittagong Veterinary & Animal Sciences University (CVASU)",
    "Bangabandhu Sheikh Mujibur Rahman Agricultural University (BSMRAU)",
    "Bangabandhu Sheikh Mujib Medical University (BSMMU)",
    "National University (NU)",
    "Bangladesh Open University (BOU)",
    "Islamic Arabic University",
    "Bangabandhu Sheikh Mujibur Rahman Science & Technology University (BSMRSTU)",
    "Bangabandhu Sheikh Mujibur Rahman Maritime University (BSMRMU)",
    "Bangabandhu Sheikh Mujibur Rahman Aviation & Aerospace University (BSMRAAU)",
    "Rabindra University, Bangladesh",
    "Sheikh Hasina University",
    "Khulna Agricultural University (KAU)",
    "Rangamati Science & Technology University (RMSTU)",
    "Chandpur Science & Technology University (CSTU)",
    "Hobiganj Agricultural University (HAU)",
    "Kurigram Agricultural University",
    "Sunamganj Science & Technology University",

    # --- Reputed Private Universities (Bangladesh) ---
    "North South University (NSU)",
    "BRAC University (BRACU)",
    "Ahsanullah University of Science & Technology (AUST)",
    "East West University (EWU)",
    "Independent University, Bangladesh (IUB)",
    "American International University-Bangladesh (AIUB)",
    "United International University (UIU)",
    "Daffodil International University (DIU)",
    "University of Asia Pacific (UAP)",
    "University of Liberal Arts Bangladesh (ULAB)",
    "Islamic University of Technology (IUT)",
    "International University of Business Agriculture & Technology (IUBAT)",
    "Stamford University Bangladesh",
    "Southeast University",
    "Green University of Bangladesh",
    "Bangladesh University of Business & Technology (BUBT)",
    "State University of Bangladesh (SUB)",
    "International Islamic University Chittagong (IIUC)",
    "Metropolitan University, Sylhet",
    "Leading University, Sylhet",
    "East Delta University (EDU)",
    "Asian University for Women (AUW)",
    "Northern University Bangladesh",
    "Uttara University",
    "World University of Bangladesh",
    "Primeasia University",
    "Manarat International University",
    "BGMEA University of Fashion & Technology (BUFT)",
    "Port City International University",
    "Premier University, Chittagong",
    "Varendra University",
    "Bangladesh Army University of Science & Technology (BAUST)",
    "Bangladesh Army University of Engineering & Technology (BAUET)",
    "Others",
]

DESIGNATION_OPTIONS = [
    "Professor",
    "Associate Professor",
    "Assistant Professor",
    "Lecturer",
    "Others",
]

COMMON_DEPARTMENTS = [
    "Electrical and Electronic Engineering",
    "Computer Science and Engineering",
    "Chemical Engineering",
    "Civil Engineering",
    "Industrial and Production Engineering",
    "Mechanical Engineering",
    "Software Engineering",
    "Physics",
    "Chemistry",
    "Mathematics",
    "Business Administration",
    "Economics",
    "English",
    "Others",
]

# -----------------------------------------------------------------------------
# 1. SESSION STATE & INITIALIZATION
# -----------------------------------------------------------------------------
if "step" not in st.session_state:
    st.session_state.step = "form"

if "form_version" not in st.session_state:
    st.session_state.form_version = 0

DEFAULT_DATA = {
    "institution_name": "",
    "custom_inst_name": "",
    "doc_type_choice": "Lab Report",
    "custom_doc_type": "",
    "work_type": "Individual",
    "exp_no": "",
    "title": "",
    "course_title": "",
    "course_code": "",
    "student_name": "",
    "reg_no": "",
    "group_no": "",
    "num_students": 2,
    "teacher_name": "",
    "designation": "Assistant Professor",
    "custom_designation": "",
    "department": "Electrical and Electronic Engineering",
    "custom_dept": "",
    "sub_date": "",
    "logo_bytes": None,
    "logo_ext": "png",
}

if "form_data" not in st.session_state:
    st.session_state.form_data = DEFAULT_DATA.copy()


def clear_all_form_fields():
    st.session_state.form_data = DEFAULT_DATA.copy()
    st.session_state.form_version += 1


def prepare_image(image_bytes):
    if not image_bytes:
        return None
    try:
        img = Image.open(io.BytesIO(image_bytes))
        return ImageReader(img)
    except Exception:
        return None


def get_final_institution_name(data):
    inst = data.get("institution_name", "")
    if inst == "Others":
        return data.get("custom_inst_name", "").strip()
    return inst


# -----------------------------------------------------------------------------
# 2. WARNING MODAL FOR BLANK FIELDS
# -----------------------------------------------------------------------------
@st.dialog("⚠️ Missing Information Warning")
def show_missing_fields_warning(missing_fields):
    st.write(
        "The following fields are currently **blank**. If you proceed, they will remain empty on the generated documents so you can write them in by hand."
    )
    for field in missing_fields:
        st.markdown(f"- **{field}**")

    st.write("Would you like to continue to the preview or edit your details?")

    col1, col2 = st.columns(2)
    with col1:
        if st.button("Proceed Anyway", type="primary", use_container_width=True):
            st.session_state.step = "preview"
            st.rerun()
    with col2:
        if st.button("Go Back & Fill", use_container_width=True):
            st.rerun()


# -----------------------------------------------------------------------------
# 3. PDF GENERATION LOGIC (WITH AUTOMATIC MULTI-LINE TITLE WRAPPING)
# -----------------------------------------------------------------------------
def generate_pdf():
    data = st.session_state.form_data
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    # Margin & printable bounds definitions
    margin = 50
    printable_width = width - (2 * margin)

    # --- 1. Document Type Header ---
    doc_type_text = data.get("doc_type_choice", "Lab Report")
    if doc_type_text == "Others":
        doc_type_text = data.get("custom_doc_type", "Cover Page")

    current_y = height - 80
    c.setFont("Times-Bold", 22)
    c.drawCentredString(width / 2, current_y, doc_type_text.upper())

    # --- 2. Top Horizontal Line ---
    current_y -= 25
    c.setLineWidth(1)
    c.line(margin, current_y, width - margin, current_y)

    # --- 3. Exp / Assignment / Project Number & Title ---
    exp_num = data.get("exp_no", "").strip()
    if exp_num:
        current_y -= 25
        c.setFont("Times-Roman", 14)
        if doc_type_text == "Lab Report":
            label_str = f"Part of Exp : {exp_num}"
        elif doc_type_text == "Assignment":
            label_str = f"Assignment No : {exp_num}"
        elif doc_type_text == "Project Report":
            label_str = f"Project No : {exp_num}"
        elif doc_type_text == "Project Proposal":
            label_str = f"Proposal No : {exp_num}"
        else:
            label_str = f"No : {exp_num}"

        c.drawCentredString(width / 2, current_y, label_str)

    # --- MULTI-LINE WRAPPED TITLE IMPLEMENTATION ---
    title_text = data.get("title", "").strip()
    if title_text:
        title_style = ParagraphStyle(
            name="TitleStyle",
            fontName="Times-Bold",
            fontSize=16,
            leading=20,
            alignment=1,  # Center Alignment
        )
        p_title = Paragraph(title_text, title_style)
        w, h = p_title.wrap(printable_width, height)

        current_y -= (h + 15)
        p_title.drawOn(c, margin, current_y)
        current_y -= 10
    else:
        current_y -= 15

    # --- 4. Bottom Horizontal Line ---
    c.line(margin, current_y, width - margin, current_y)

    # --- 5. Optional Logo ---
    logo_img = prepare_image(data.get("logo_bytes"))
    if logo_img:
        current_y -= 100
        logo_w, logo_h = 90, 90
        c.drawImage(
            logo_img,
            (width - logo_w) / 2,
            current_y,
            width=logo_w,
            height=logo_h,
            mask="auto",
        )
        current_y -= 20
    else:
        current_y -= 40

    # --- 6. Department Name & Course Info ---
    dept_val = data.get("department", "")
    if dept_val == "Others":
        dept_val = data.get("custom_dept", "")
    dept_str = (
        f"Department of {dept_val}"
        if dept_val and not dept_val.lower().startswith("department")
        else dept_val
    )

    if dept_str:
        c.setFont("Times-Bold", 15)
        c.drawCentredString(width / 2, current_y, dept_str)
        current_y -= 25

    if data.get("course_title"):
        c.setFont("Times-Roman", 13)
        c.drawCentredString(width / 2, current_y, f"Course Title: {data['course_title']}")
        current_y -= 20

    if data.get("course_code"):
        c.setFont("Times-Roman", 13)
        c.drawCentredString(width / 2, current_y, f"Course Code: {data['course_code']}")
        current_y -= 20

    # --- 7. Submitted By Section ---
    current_y -= 35
    c.setFont("Times-Bold", 13)
    c.drawCentredString(width / 2, current_y, "Submitted By:")
    c.line(width / 2 - 45, current_y - 3, width / 2 + 45, current_y - 3)

    c.setFont("Times-Roman", 12)
    current_y -= 25

    if data.get("work_type") == "Individual":
        if data.get("student_name"):
            c.drawCentredString(width / 2, current_y, data["student_name"])
            current_y -= 20
        if data.get("reg_no"):
            c.drawCentredString(width / 2, current_y, f"Registration No: {data['reg_no']}")
            current_y -= 20
    else:
        for i in range(int(data.get("num_students", 2))):
            s_name = data.get(f"sname_{i}", "")
            s_reg = data.get(f"sreg_{i}", "")
            if s_name:
                info = f"{s_name} (Reg: {s_reg})" if s_reg else s_name
                c.drawCentredString(width / 2, current_y, info)
                current_y -= 18

        if data.get("group_no"):
            current_y -= 5
            c.drawCentredString(width / 2, current_y, f"Group No : {data['group_no']}")
            current_y -= 20

    # --- 8. Submitted To Section ---
    current_y -= 25
    c.setFont("Times-Bold", 13)
    c.drawCentredString(width / 2, current_y, "Submitted To:")
    c.line(width / 2 - 45, current_y - 3, width / 2 + 45, current_y - 3)

    c.setFont("Times-Roman", 12)
    current_y -= 25

    if data.get("teacher_name"):
        c.drawCentredString(width / 2, current_y, data["teacher_name"])
        current_y -= 20

    desig_val = data.get("designation", "")
    if desig_val == "Others":
        desig_val = data.get("custom_designation", "")
    if desig_val:
        c.drawCentredString(width / 2, current_y, desig_val)
        current_y -= 20

    if dept_str:
        c.drawCentredString(width / 2, current_y, dept_str)
        current_y -= 20

    inst_name = get_final_institution_name(data)
    if inst_name:
        c.drawCentredString(width / 2, current_y, inst_name)

    # --- 9. Date of Submission ---
    sub_date = data.get("sub_date", "").strip()
    c.setFont("Times-Roman", 12)
    if sub_date:
        c.drawCentredString(width / 2, 50, f"Date of Submission: {sub_date}")
    else:
        c.drawCentredString(width / 2, 50, "Date of Submission: ....................................")

    c.showPage()
    c.save()
    buffer.seek(0)
    return buffer


# -----------------------------------------------------------------------------
# 4. DOCX GENERATION LOGIC (NATIVE MS WORD)
# -----------------------------------------------------------------------------
def generate_docx():
    data = st.session_state.form_data
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    def add_centered_p(text="", bold=False, size=12, space_after=6, underline=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(space_after)
        if text:
            run = p.add_run(text)
            run.bold = bold
            run.underline = underline
            run.font.name = "Times New Roman"
            run.font.size = Pt(size)
        return p

    # --- 1. Document Type Header ---
    doc_type_text = data.get("doc_type_choice", "Lab Report")
    if doc_type_text == "Others":
        doc_type_text = data.get("custom_doc_type", "Cover Page")

    add_centered_p(doc_type_text.upper(), bold=True, size=22, space_after=12)

    # --- 2. Exp / Assignment / Proposal Number & Title ---
    exp_num = data.get("exp_no", "").strip()
    if exp_num:
        if doc_type_text == "Lab Report":
            label_str = f"Part of Exp : {exp_num}"
        elif doc_type_text == "Assignment":
            label_str = f"Assignment No : {exp_num}"
        elif doc_type_text == "Project Report":
            label_str = f"Project No : {exp_num}"
        elif doc_type_text == "Project Proposal":
            label_str = f"Proposal No : {exp_num}"
        else:
            label_str = f"No : {exp_num}"

        add_centered_p(label_str, size=14, space_after=6)

    if data.get("title"):
        add_centered_p(data["title"], bold=True, size=16, space_after=18)

    # --- 3. Optional Logo ---
    if data.get("logo_bytes"):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(12)
        p_img.paragraph_format.space_after = Pt(18)
        image_stream = io.BytesIO(data["logo_bytes"])
        p_img.add_run().add_picture(image_stream, width=Inches(1.2))

    # --- 4. Department Name & Course Info ---
    dept_val = data.get("department", "")
    if dept_val == "Others":
        dept_val = data.get("custom_dept", "")
    dept_str = (
        f"Department of {dept_val}"
        if dept_val and not dept_val.lower().startswith("department")
        else dept_val
    )

    if dept_str:
        add_centered_p(dept_str, bold=True, size=15, space_after=12)

    if data.get("course_title"):
        add_centered_p(f"Course Title: {data['course_title']}", size=13, space_after=4)

    if data.get("course_code"):
        add_centered_p(f"Course Code: {data['course_code']}", size=13, space_after=24)

    # --- 5. Submitted By Section ---
    add_centered_p("Submitted By:", bold=True, size=13, space_after=6, underline=True)

    if data.get("work_type") == "Individual":
        if data.get("student_name"):
            add_centered_p(data["student_name"], size=12, space_after=4)
        if data.get("reg_no"):
            add_centered_p(f"Registration No: {data['reg_no']}", size=12, space_after=20)
    else:
        for i in range(int(data.get("num_students", 2))):
            s_name = data.get(f"sname_{i}", "")
            s_reg = data.get(f"sreg_{i}", "")
            if s_name:
                info = f"{s_name} (Reg: {s_reg})" if s_reg else s_name
                add_centered_p(info, size=12, space_after=4)

        if data.get("group_no"):
            add_centered_p(f"Group No : {data['group_no']}", size=12, space_after=20)

    # --- 6. Submitted To Section ---
    add_centered_p("Submitted To:", bold=True, size=13, space_after=6, underline=True)

    if data.get("teacher_name"):
        add_centered_p(data["teacher_name"], size=12, space_after=4)

    desig_val = data.get("designation", "")
    if desig_val == "Others":
        desig_val = data.get("custom_designation", "")
    if desig_val:
        add_centered_p(desig_val, size=12, space_after=4)

    if dept_str:
        add_centered_p(dept_str, size=12, space_after=4)

    inst_name = get_final_institution_name(data)
    if inst_name:
        add_centered_p(inst_name, size=12, space_after=30)

    # --- 7. Date of Submission ---
    sub_date = data.get("sub_date", "").strip()
    date_text = (
        f"Date of Submission: {sub_date}"
        if sub_date
        else "Date of Submission: ...................................."
    )
    add_centered_p(date_text, size=12, space_after=0)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# -----------------------------------------------------------------------------
# 5. LATEX GENERATION & ZIP BUNDLING LOGIC
# -----------------------------------------------------------------------------
def generate_tex():
    data = st.session_state.form_data

    doc_type = data.get("doc_type_choice", "Lab Report")
    if doc_type == "Others":
        doc_type = data.get("custom_doc_type", "Cover Page")

    exp_num = data.get("exp_no", "").strip()
    if exp_num:
        if doc_type == "Lab Report":
            num_str = f"Part of Exp : {exp_num}"
        elif doc_type == "Assignment":
            num_str = f"Assignment No : {exp_num}"
        elif doc_type == "Project Report":
            num_str = f"Project No : {exp_num}"
        elif doc_type == "Project Proposal":
            num_str = f"Proposal No : {exp_num}"
        else:
            num_str = f"No : {exp_num}"
    else:
        num_str = ""

    student_block = ""
    if data.get("work_type") == "Individual":
        if data.get("student_name"):
            student_block += f"{data['student_name']}\\\\\n"
        if data.get("reg_no"):
            student_block += f"Registration No: {data['reg_no']}\\\\\n"
    else:
        for i in range(int(data.get("num_students", 2))):
            s_name = data.get(f"sname_{i}", "")
            s_reg = data.get(f"sreg_{i}", "")
            if s_name:
                info = f"{s_name} (Reg: {s_reg})" if s_reg else s_name
                student_block += f"{info}\\\\\n"
        if data.get("group_no"):
            student_block += f"Group No : {data['group_no']}\\\\\n"

    desig_val = data.get("designation", "")
    if desig_val == "Others":
        desig_val = data.get("custom_designation", "")

    dept_val = data.get("department", "")
    if dept_val == "Others":
        dept_val = data.get("custom_dept", "")
    dept_str = (
        f"Department of {dept_val}"
        if dept_val and not dept_val.lower().startswith("department")
        else dept_val
    )

    # Escape LaTeX special character &
    dept_str_tex = dept_str.replace("&", "\\&")
    inst_name_tex = get_final_institution_name(data).replace("&", "\\&")

    sub_date = data.get("sub_date", "").strip()
    date_str = (
        f"Date of Submission: {sub_date}"
        if sub_date
        else "Date of Submission: \\dots\\dots\\dots\\dots\\dots\\dots\\dots\\dots\\dots\\dots\\dots\\dots\\dots\\dots"
    )

    logo_filename = f"logo.{data.get('logo_ext', 'png')}"
    logo_code = (
        f"\\IfFileExists{{{logo_filename}}}{{\n    \\includegraphics[width=1.2in]{{{logo_filename}}} \\\\[1.5em]\n}}{{\n    \\vspace{{1.2in}}\n}}"
        if data.get("logo_bytes")
        else "\\vspace{1.5em}"
    )

    tex_code = f"""\\documentclass[12pt,a4paper]{{article}}
\\usepackage[margin=1in]{{geometry}}
\\usepackage{{graphicx}}
\\usepackage{{setspace}}
\\pagestyle{{empty}}

\\begin{{document}}
\\begin{{center}}

{{\\Huge \\textbf{{{doc_type.upper()}}}}} \\\\[1.5em]
\\hrule height 1pt
\\vspace{{1em}}

\\Large {num_str} \\\\[0.8em]
{{\\Large \\textbf{{{data.get("title", "")}}}}} \\\\[1em]

\\hrule height 1pt
\\vspace{{2em}}

{logo_code}

{{\\large \\textbf{{{dept_str_tex}}}}} \\\\[0.8em]
Course Title: {data.get("course_title", "")} \\\\[0.5em]
Course Code: {data.get("course_code", "")} \\\\[2em]

{{\\large \\underline{{\\textbf{{Submitted By:}}}}}} \\\\[0.8em]
{student_block}
\\vspace{{1.5em}}

{{\\large \\underline{{\\textbf{{Submitted To:}}}}}} \\\\[0.8em]
{data.get("teacher_name", "")} \\\\[0.4em]
{desig_val} \\\\[0.4em]
{dept_str_tex} \\\\[0.4em]
{inst_name_tex} \\\\[3em]

\\vfill
{date_str}

\\end{{center}}
\\end{{document}}
"""
    return tex_code


def generate_tex_zip(tex_code):
    data = st.session_state.form_data
    zip_buffer = io.BytesIO()

    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        zip_file.writestr("Cover_Page.tex", tex_code)

        if data.get("logo_bytes"):
            ext = data.get("logo_ext", "png")
            zip_file.writestr(f"logo.{ext}", data["logo_bytes"])

    zip_buffer.seek(0)
    return zip_buffer


def get_pdf_preview(pdf_bytes):
    pdf = pdfium.PdfDocument(pdf_bytes)
    page = pdf[0]
    image = page.render(scale=2).to_pil()
    return image


# -----------------------------------------------------------------------------
# 6. STREAMLIT UI LAYOUT
# -----------------------------------------------------------------------------
if st.session_state.step == "form":
    st.title("📄 Cover Page Generator")

    doc_options = ["Lab Report", "Assignment", "Project Report", "Project Proposal", "Others"]
    work_options = ["Individual", "Group Work"]

    v = st.session_state.form_version
    d = st.session_state.form_data

    # Institution Name & Optional Logo Uploader
    st.subheader("Institution Information")

    curr_inst = d.get("institution_name", "")
    inst_idx = (
        BANGLADESH_UNIVERSITIES.index(curr_inst)
        if curr_inst in BANGLADESH_UNIVERSITIES
        else 0
    )

    selected_inst = st.selectbox(
        "Institution Name *",
        options=["Select a University..."] + BANGLADESH_UNIVERSITIES,
        index=inst_idx + 1 if curr_inst in BANGLADESH_UNIVERSITIES else 0,
        key=f"inst_select_{v}",
        help="Select your university from the list or choose 'Others' to type a custom name.",
    )

    if selected_inst == "Select a University...":
        d["institution_name"] = ""
    else:
        d["institution_name"] = selected_inst

    if d["institution_name"] == "Others":
        d["custom_inst_name"] = st.text_input(
            "Specify Institution Name *",
            value=d.get("custom_inst_name", ""),
            key=f"custom_inst_{v}",
            placeholder="Type your university name here...",
        )

    uploaded_logo = st.file_uploader(
        "Upload Institution Logo (Optional)",
        type=["png", "jpg", "jpeg", "webp", "bmp", "svg"],
        key=f"logo_upload_{v}",
    )
    if uploaded_logo is not None:
        d["logo_bytes"] = uploaded_logo.getvalue()
        d["logo_ext"] = uploaded_logo.name.split(".")[-1].lower()

    # Document Details
    st.subheader("Document Details")
    current_choice = d.get("doc_type_choice", "Lab Report")
    idx_doc = doc_options.index(current_choice) if current_choice in doc_options else 0

    d["doc_type_choice"] = st.selectbox(
        "Select Cover Page Type",
        doc_options,
        index=idx_doc,
        key=f"doc_type_{v}",
    )

    if d["doc_type_choice"] == "Others":
        d["custom_doc_type"] = st.text_input(
            "Specify Cover Page Type",
            value=d.get("custom_doc_type", ""),
            key=f"custom_doc_{v}",
        )

    exp_field_label = "Experiment No (e.g., 02)"
    if d["doc_type_choice"] == "Assignment":
        exp_field_label = "Assignment No (e.g., 01)"
    elif d["doc_type_choice"] == "Project Report":
        exp_field_label = "Project No (e.g., 01)"
    elif d["doc_type_choice"] == "Project Proposal":
        exp_field_label = "Proposal No (e.g., 01)"
    elif d["doc_type_choice"] == "Others":
        exp_field_label = "Number / ID (e.g., 01)"

    d["exp_no"] = st.text_input(
        exp_field_label,
        value=d.get("exp_no", ""),
        key=f"exp_no_{v}",
    )
    d["title"] = st.text_area(
        "Title",
        value=d.get("title", ""),
        key=f"title_{v}",
        help="Long titles will automatically wrap into multiple centered lines on the PDF output.",
    )
    d["course_title"] = st.text_input(
        "Course Title",
        value=d.get("course_title", ""),
        key=f"course_title_{v}",
    )
    d["course_code"] = st.text_input(
        "Course Code",
        value=d.get("course_code", ""),
        key=f"course_code_{v}",
    )

    # Student Information
    st.subheader("Student Details")
    d["work_type"] = st.radio(
        "Select Submission Type",
        work_options,
        index=work_options.index(d.get("work_type", "Individual")),
        horizontal=True,
        key=f"work_type_{v}",
    )

    if d["work_type"] == "Individual":
        d["student_name"] = st.text_input(
            "Student Name",
            value=d.get("student_name", ""),
            key=f"student_name_{v}",
        )
        d["reg_no"] = st.text_input(
            "Registration No",
            value=d.get("reg_no", ""),
            key=f"reg_no_{v}",
        )
    else:
        d["group_no"] = st.text_input(
            "Group No",
            value=d.get("group_no", ""),
            key=f"group_no_{v}",
        )
        d["num_students"] = st.number_input(
            "Number of Students",
            min_value=2,
            max_value=10,
            value=int(d.get("num_students", 2)),
            key=f"num_students_{v}",
        )

        for i in range(int(d["num_students"])):
            col1, col2 = st.columns(2)
            d[f"sname_{i}"] = col1.text_input(
                f"Student {i + 1} Name",
                value=d.get(f"sname_{i}", ""),
                key=f"sname_{i}_{v}",
            )
            d[f"sreg_{i}"] = col2.text_input(
                f"Student {i + 1} Reg No",
                value=d.get(f"sreg_{i}", ""),
                key=f"sreg_{i}_{v}",
            )

    # Teacher Details Section
    st.subheader("Teacher Details")

    d["teacher_name"] = st.text_input(
        "Teacher Name",
        value=d.get("teacher_name", ""),
        key=f"tname_{v}",
    )

    curr_desig = d.get("designation", "Assistant Professor")
    desig_idx = (
        DESIGNATION_OPTIONS.index(curr_desig)
        if curr_desig in DESIGNATION_OPTIONS
        else 4
    )

    selected_desig = st.selectbox(
        "Teacher Designation",
        DESIGNATION_OPTIONS,
        index=desig_idx,
        key=f"desig_select_{v}",
    )

    if selected_desig == "Others":
        d["designation"] = "Others"
        d["custom_designation"] = st.text_input(
            "Specify Designation",
            value=d.get("custom_designation", ""),
            key=f"custom_desig_{v}",
        )
    else:
        d["designation"] = selected_desig

    curr_dept = d.get("department", "Electrical and Electronic Engineering")
    dept_idx = (
        COMMON_DEPARTMENTS.index(curr_dept)
        if curr_dept in COMMON_DEPARTMENTS
        else (len(COMMON_DEPARTMENTS) - 1)
    )

    selected_dept = st.selectbox(
        "Department Name",
        COMMON_DEPARTMENTS,
        index=dept_idx,
        key=f"dept_select_{v}",
    )

    if selected_dept == "Others":
        d["department"] = "Others"
        d["custom_dept"] = st.text_input(
            "Specify Department Name",
            value=d.get("custom_dept", ""),
            key=f"custom_dept_{v}",
        )
    else:
        d["department"] = selected_dept

    d["sub_date"] = st.text_input(
        "Date of Submission (Leave blank for printable dotted line)",
        value=d.get("sub_date", ""),
        key=f"sub_date_{v}",
    )

    st.session_state.form_data = d

    action_col1, action_col2 = st.columns(2)

    with action_col1:
        if st.button("Generate & Preview", type="primary", use_container_width=True):
            missing_fields = []
            final_inst = get_final_institution_name(d)
            if not final_inst:
                missing_fields.append("Institution Name")
            if not d.get("exp_no", "").strip():
                missing_fields.append(exp_field_label)
            if not d.get("title", "").strip():
                missing_fields.append("Title")
            if not d.get("course_title", "").strip():
                missing_fields.append("Course Title")
            if not d.get("course_code", "").strip():
                missing_fields.append("Course Code")

            if d["work_type"] == "Individual":
                if not d.get("student_name", "").strip():
                    missing_fields.append("Student Name")
                if not d.get("reg_no", "").strip():
                    missing_fields.append("Registration No")
            else:
                if not d.get("group_no", "").strip():
                    missing_fields.append("Group No")

            if not d.get("teacher_name", "").strip():
                missing_fields.append("Teacher Name")

            if d.get("department") == "Others" and not d.get("custom_dept", "").strip():
                missing_fields.append("Specified Department")

            if missing_fields:
                show_missing_fields_warning(missing_fields)
            else:
                st.session_state.step = "preview"
                st.rerun()

    with action_col2:
        if st.button("🗑️ Clear Form", use_container_width=True):
            clear_all_form_fields()
            st.rerun()

elif st.session_state.step == "preview":
    st.title("📄 Cover Page Generator")
    st.subheader("Preview Your Cover Page")

    pdf_buffer = generate_pdf()
    pdf_bytes = pdf_buffer.getvalue()

    docx_buffer = generate_docx()
    docx_bytes = docx_buffer.getvalue()

    tex_code = generate_tex()
    tex_zip_buffer = generate_tex_zip(tex_code)
    tex_zip_bytes = tex_zip_buffer.getvalue()

    preview_img = get_pdf_preview(pdf_bytes)
    st.image(preview_img, use_container_width=True)

    btn_col1, btn_col2, btn_col3, btn_col4 = st.columns([1, 1, 1, 1.2])

    with btn_col1:
        if st.button("✏️ Edit Details", use_container_width=True):
            st.session_state.step = "form"
            st.rerun()

    with btn_col2:
        st.download_button(
            label="📥 PDF",
            data=pdf_bytes,
            file_name="Cover_Page.pdf",
            mime="application/pdf",
            type="primary",
            use_container_width=True,
        )

    with btn_col3:
        st.download_button(
            label="📝 DOCX",
            data=docx_bytes,
            file_name="Cover_Page.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

    with btn_col4:
        st.download_button(
            label="📦 LaTeX Bundle (.zip)",
            data=tex_zip_bytes,
            file_name="LaTeX_Cover_Page.zip",
            mime="application/zip",
            use_container_width=True,
        )
